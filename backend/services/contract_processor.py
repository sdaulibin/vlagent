#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
合同比对处理服务

功能：
1. 文档转图片
2. 使用 AI 提取文本内容
3. 文本差异比对
4. 生成差异报告
"""

import os
import json
from pathlib import Path
from pdf2image import convert_from_path
from src.config import UPLOAD_DIR, DOWNLOAD_DIR
from services.core.request_ai import request_qwen35
from src.json_repair import fix_json
from services.pdf_processor import load_schema

# 合同比对输出目录
CONTRACT_OUTPUT_DIR = os.path.join(DOWNLOAD_DIR, "contracts")
os.makedirs(CONTRACT_OUTPUT_DIR, exist_ok=True)


def extract_text_from_image(image_path: str) -> str:
    """
    使用 AI 从图片中提取文本内容
    
    Args:
        image_path: 图片文件路径
        
    Returns:
        str: 提取的文本内容
    """
    prompt = """
    你是一个专业的文档文本提取专家。请仔细阅读图片中的所有文字内容，并完整提取出来。
    
    要求：
    1. 保持原文的段落结构和层级关系
    2. 保留标题、编号、列表等格式
    3. 如果有表格，用文字描述表格内容
    4. 忽略页眉、页脚、页码等非正文内容
    5. 输出纯文本，不需要 JSON 格式
    6. 如果图片中没有文字，直接返回空
    
    直接输出提取的文本内容，不需要任何解释。
    """
    
    result = request_qwen35(
        question=prompt,
        show_request=False,
        file_base=image_path
    )
    
    # 清理无效的 OCR 结果
    if result:
        # 移除 markdown 代码块
        import re
        result = re.sub(r'```\w*\n?', '', result)
        result = re.sub(r'```', '', result)
        
        # 移除 HTML 标签和无文字提示
        result = re.sub(r'<[^>]+>', '', result)
        result = result.replace('图中无文字信息', '')
        result = result.replace('图中无文字', '')
        
        result = result.strip()
    
    return result


def convert_to_images(file_path: str, output_folder: str) -> list:
    """
    将文档转换为图片列表
    
    Args:
        file_path: 文档路径 (PDF)
        output_folder: 输出文件夹
        
    Returns:
        list: 图片路径列表
    """
    os.makedirs(output_folder, exist_ok=True)
    
    # PDF 文件
    if file_path.lower().endswith('.pdf'):
        pages = convert_from_path(file_path, dpi=200)
        image_paths = []
        
        filename = Path(file_path).stem
        for i, page in enumerate(pages):
            image_path = os.path.join(output_folder, f"{filename}_page_{i+1:03d}.png")
            page.save(image_path, 'PNG')
            image_paths.append(image_path)
        
        return image_paths
    
    # 图片文件直接返回
    elif file_path.lower().endswith(('.jpg', '.jpeg', '.png')):
        return [file_path]
    
    # DOCX 文件返回空列表（使用文本提取方式）
    elif file_path.lower().endswith(('.docx', '.doc')):
        return []
    
    else:
        raise ValueError(f"Unsupported file format: {file_path}")


def extract_text_from_docx(file_path: str) -> str:
    """
    从 DOCX 文件中提取文本
    
    Args:
        file_path: DOCX 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        from docx import Document
        import shutil
        import tempfile
        import uuid
        
        # 处理中文文件名编码问题：复制到临时目录
        tmp_dir = tempfile.mkdtemp()
        tmp_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex}.docx")
        
        try:
            # 使用二进制读写，避免编码问题
            with open(file_path, 'rb') as src:
                content = src.read()
            with open(tmp_path, 'wb') as dst:
                dst.write(content)
            
            doc = Document(tmp_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return '\n\n'.join(paragraphs)
        finally:
            # 清理临时目录
            shutil.rmtree(tmp_dir, ignore_errors=True)
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")


def extract_text_from_doc(file_path: str) -> str:
    """
    从旧版 DOC 文件中提取文本 (使用 antiword 或 textract)
    
    Args:
        file_path: DOC 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    import subprocess
    
    # 尝试使用 antiword
    try:
        result = subprocess.run(
            ['antiword', file_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            return result.stdout
    except FileNotFoundError:
        print("antiword 未安装，尝试其他方法...")
    except subprocess.TimeoutExpired:
        print("antiword 超时")
    except Exception as e:
        print(f"antiword 失败: {e}")
    
    # 尝试使用 LibreOffice 转换
    try:
        import tempfile
        with tempfile.TemporaryDirectory() as temp_dir:
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'txt:Text', '--outdir', temp_dir, file_path],
                capture_output=True,
                text=True,
                timeout=120
            )
            if result.returncode == 0:
                # 找到转换后的 txt 文件
                txt_file = os.path.join(temp_dir, Path(file_path).stem + '.txt')
                if os.path.exists(txt_file):
                    with open(txt_file, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
    except FileNotFoundError:
        print("LibreOffice 未安装")
    except subprocess.TimeoutExpired:
        print("LibreOffice 转换超时")
    except Exception as e:
        print(f"LibreOffice 转换失败: {e}")
    
    # 最后尝试：读取为二进制并提取可读文本
    print("使用基础文本提取...")
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            # 尝试提取可读文本
            text = content.decode('utf-8', errors='ignore')
            # 过滤掉不可打印字符
            import re
            text = re.sub(r'[^\u4e00-\u9fff\u0020-\u007e\n\r\t]', '', text)
            # 清理多余空行
            text = re.sub(r'\n{3,}', '\n\n', text)
            return text.strip()
    except Exception as e:
        raise ValueError(f"无法提取 DOC 文件内容: {e}")


def extract_document_content(file_path: str) -> str:
    """
    提取文档的全部文本内容
    
    Args:
        file_path: 文档路径
        
    Returns:
        str: 文档文本内容
    """
    # DOCX 文件直接提取文本
    if file_path.lower().endswith('.docx'):
        print(f"提取 DOCX 文本: {file_path}")
        return extract_text_from_docx(file_path)
    
    # 旧版 DOC 文件
    if file_path.lower().endswith('.doc'):
        print(f"提取 DOC 文本: {file_path}")
        return extract_text_from_doc(file_path)
    
    # 其他格式转图片后 OCR
    filename = Path(file_path).stem
    output_folder = os.path.join(CONTRACT_OUTPUT_DIR, f"task_{filename}")
    
    # 转换为图片
    image_paths = convert_to_images(file_path, output_folder)
    
    if not image_paths:
        raise ValueError(f"无法处理文件: {file_path}")
    
    # 提取每页文本
    all_text = []
    for image_path in image_paths:
        text = extract_text_from_image(image_path)
        all_text.append(text)
    
    return "\n\n".join(all_text)


def _compare_chunk_with_difflib(chunk_a: str, chunk_b: str) -> list:
    """
    使用 difflib 程序化比对文档片段，精确找出所有差异
    """
    import difflib
    
    if not chunk_a.strip() and not chunk_b.strip():
        return []
    
    # 按行分割
    lines_a = chunk_a.split('\n')
    lines_b = chunk_b.split('\n')
    
    diffs = []
    matcher = difflib.SequenceMatcher(None, lines_a, lines_b)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        original_lines = lines_a[i1:i2]
        comparison_lines = lines_b[j1:j2]
        
        original_text = '\n'.join(original_lines).strip()
        comparison_text = '\n'.join(comparison_lines).strip()
        
        # 跳过纯空白差异
        if not original_text and not comparison_text:
            continue
        
        if tag == 'replace':
            diffs.append({
                "type": "modified",
                "original": original_text,
                "comparison": comparison_text,
                "location": f"第{i1+1}行附近"
            })
        elif tag == 'delete':
            diffs.append({
                "type": "deleted",
                "original": original_text,
                "comparison": "",
                "location": f"第{i1+1}行附近"
            })
        elif tag == 'insert':
            diffs.append({
                "type": "added",
                "original": "",
                "comparison": comparison_text,
                "location": f"第{j1+1}行附近"
            })
    
    # 如果按行比对没有发现差异，进行字符级比对
    if not diffs and chunk_a.strip() != chunk_b.strip():
        char_matcher = difflib.SequenceMatcher(None, chunk_a, chunk_b)
        for tag, i1, i2, j1, j2 in char_matcher.get_opcodes():
            if tag == 'equal':
                continue
            
            original_chars = chunk_a[i1:i2]
            comparison_chars = chunk_b[j1:j2]
            
            # 获取上下文
            context_start = max(0, i1 - 30)
            context_end = min(len(chunk_a), i2 + 30)
            context = chunk_a[context_start:context_end].replace('\n', ' ')
            
            if tag == 'replace':
                diffs.append({
                    "type": "modified",
                    "original": original_chars,
                    "comparison": comparison_chars,
                    "location": f"...{context}..."
                })
            elif tag == 'delete':
                diffs.append({
                    "type": "deleted",
                    "original": original_chars,
                    "comparison": "",
                    "location": f"...{context}..."
                })
            elif tag == 'insert':
                diffs.append({
                    "type": "added",
                    "original": "",
                    "comparison": comparison_chars,
                    "location": f"新增内容"
                })
    
    return diffs


def compare_texts(text_a: str, text_b: str) -> list:
    """
    智能分块比对长文档
    """
    # 限制单个分块大小 (字符数)
    # 降低到 3000 以避免本地模型卡死
    CHUNK_SIZE = 3000
    
    # 如果文档较小，直接比对
    if len(text_a) < CHUNK_SIZE and len(text_b) < CHUNK_SIZE:
        print(f"文档较短 ({len(text_a)} chars)，直接比对")
        return _compare_chunk_with_difflib(text_a, text_b)
    
    print(f"文档较长 ({len(text_a)} chars)，开始分块比对...")
    import difflib
    
    # 使用 difflib 找到匹配块，以此作为分割点
    s = difflib.SequenceMatcher(None, text_a, text_b)
    matching_blocks = s.get_matching_blocks()
    
    all_diffs = []
    last_a = 0
    last_b = 0
    
    # 当前累积的待比对文本
    pending_a = []
    pending_b = []
    
    for i, block in enumerate(matching_blocks):
        a_start, b_start, size = block
        
        # 获取不匹配的部分
        unmatched_a = text_a[last_a:a_start]
        unmatched_b = text_b[last_b:b_start]
        
        # 获取匹配的部分
        matched_content = text_a[a_start:a_start+size]
        
        # 累积
        pending_a.append(unmatched_a)
        pending_b.append(unmatched_b)
        pending_a.append(matched_content)
        pending_b.append(matched_content)
        
        current_str_a = "".join(pending_a)
        current_str_b = "".join(pending_b)
        
        # 检查是否需要提交分块
        if len(current_str_a) >= CHUNK_SIZE or i == len(matching_blocks) - 1:
            if current_str_a.strip() or current_str_b.strip():
                chunk_num = len(all_diffs) + 1
                print(f"处理分块 #{chunk_num}: Size A={len(current_str_a)}, Size B={len(current_str_b)}")
                try:
                    chunk_diffs = _compare_chunk_with_difflib(current_str_a, current_str_b)
                    print(f"分块 #{chunk_num} 完成，发现 {len(chunk_diffs)} 处差异")
                    all_diffs.extend(chunk_diffs)
                except Exception as e:
                    print(f"分块 #{chunk_num} 处理出错: {e}")
            
            # 重置
            pending_a = []
            pending_b = []
        
        last_a = a_start + size
        last_b = b_start + size
        
    return all_diffs


def compare_documents(file_a_path: str, file_b_path: str) -> list:
    """
    比对两份文档
    
    Args:
        file_a_path: 原文档路径
        file_b_path: 比对文档路径
        
    Returns:
        list: 差异列表
    """
    result = compare_documents_with_content(file_a_path, file_b_path)
    return result.get("diffs", [])


def compare_documents_with_content(file_a_path: str, file_b_path: str) -> dict:
    """
    比对两份文档，返回内容和差异
    
    Args:
        file_a_path: 原文档路径
        file_b_path: 比对文档路径
        
    Returns:
        dict: 包含 content_a, content_b, diffs
    """
    print(f"开始比对文档: {file_a_path} vs {file_b_path}")
    
    # 提取文本内容
    print("提取原文档内容...")
    text_a = extract_document_content(file_a_path)
    
    print("提取比对文档内容...")
    text_b = extract_document_content(file_b_path)
    
    # 比对差异
    print("分析文档差异...")
    diffs = compare_texts(text_a, text_b)
    
    print(f"共发现 {len(diffs)} 处差异")
    
    return {
        "content_a": text_a,
        "content_b": text_b,
        "diffs": diffs
    }


if __name__ == "__main__":
    # 测试代码
    test_a = os.path.join(UPLOAD_DIR, "contracts/test_a.pdf")
    test_b = os.path.join(UPLOAD_DIR, "contracts/test_b.pdf")
    
    if os.path.exists(test_a) and os.path.exists(test_b):
        result = compare_documents(test_a, test_b)
        # 结果不再直接打印到控制台
