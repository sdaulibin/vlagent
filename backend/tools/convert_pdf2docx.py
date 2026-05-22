"""
PDF 转 DOCX 独立工具

默认使用混合模式：pdfplumber 检测表格结构 + RapidOCR 提取正确文字。
适用于字体编码异常的 PDF（如国家标准文档）。

用法:
    uv run python tools/convert_pdf2docx.py input.pdf
    uv run python tools/convert_pdf2docx.py input.pdf -o output.docx
    uv run python tools/convert_pdf2docx.py input.pdf --pages 1-3,5
    uv run python tools/convert_pdf2docx.py input.pdf --text    # 纯文本模式（快）
    uv run python tools/convert_pdf2docx.py ./pdfs/             # 批量转换
"""

import argparse
import io
import sys
from pathlib import Path

import fitz  # pymupdf
import pdfplumber
from docx import Document
from docx.shared import Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image


def _parse_pages(pages_str: str) -> list[int]:
    result = []
    for part in pages_str.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            result.extend(range(int(start) - 1, int(end)))
        else:
            result.append(int(part) - 1)
    return result


# ── OCR 引擎 ──────────────────────────────────────────────────────────

_ocr_engine = None


def _get_ocr():
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    return _ocr_engine


def _ocr_image(img: Image.Image) -> list[tuple[str, tuple]]:
    """OCR 图片，返回 [(text, bbox), ...]"""
    ocr = _get_ocr()
    result, _ = ocr(img)
    if not result:
        return []
    return [(item[1], item[0]) for item in result]


def _render_page(page, dpi: int = 300) -> Image.Image:
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat)
    return Image.open(io.BytesIO(pix.tobytes("png")))


def _ocr_region(img: Image.Image, x0, y0, x1, y1, scale: float) -> str:
    """裁剪图片区域并 OCR"""
    crop = img.crop((x0 * scale, y0 * scale, x1 * scale, y1 * scale))
    items = _ocr_image(crop)
    if not items:
        return ""
    items.sort(key=lambda it: (round(it[1][0][1] / 10), it[1][0][0]))
    lines = []
    current_line = []
    last_y = None
    for text, bbox in items:
        y = bbox[0][1]
        if last_y is not None and abs(y - last_y) > 15:
            if current_line:
                lines.append(" ".join(current_line))
            current_line = []
        current_line.append(text)
        last_y = y
    if current_line:
        lines.append(" ".join(current_line))
    return "\n".join(lines)


# ── 混合模式表格处理 ─────────────────────────────────────────────────

def _get_cell_bboxes(table) -> list[list[tuple | None]]:
    """从 pdfplumber table 获取每行每列的 bbox"""
    result = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            row_cells.append(cell)
        result.append(row_cells)
    return result


def _is_header_row(row_cells: list, ncols: int) -> bool:
    """检测是否为表头行（大部分单元格为 None 表示跨列合并）"""
    non_none = sum(1 for c in row_cells if c is not None)
    if non_none <= 1 and ncols > 1:
        return True
    return False


def _build_table(doc: Document, table, full_img: Image.Image, scale: float) -> None:
    """构建 DOCX 表格：正确的列宽 + OCR 文字 + 对齐"""
    rows_data = table.extract()
    if not rows_data:
        return

    cell_bboxes = _get_cell_bboxes(table)
    ncols = max(len(r) for r in rows_data)

    # 计算列宽（从单元格 bbox 推算）
    col_widths = _calc_col_widths(cell_bboxes, ncols, table.bbox)

    # 检测表头行（合并单元格的标题行）
    header_start = 0
    for ri, row_cells in enumerate(cell_bboxes):
        if _is_header_row(row_cells, ncols):
            # 把表头作为标题段落，不放入表格
            if row_cells[0] is not None:
                cx0, cy0, cx1, cy1 = row_cells[0]
                title_text = _ocr_region(full_img, cx0, cy0, cx1, cy1, scale)
                if title_text.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(title_text.strip())
                    run.bold = True
            header_start = ri + 1
        else:
            break

    # 剩余行构成表格
    data_rows = rows_data[header_start:]
    data_bboxes = cell_bboxes[header_start:]
    if not data_rows:
        return

    actual_ncols = max(len(r) for r in data_rows)
    t = doc.add_table(rows=len(data_rows), cols=actual_ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    for ci in range(actual_ncols):
        if ci < len(col_widths):
            t.columns[ci].width = col_widths[ci]

    for ri, row in enumerate(data_rows):
        row_bboxes = data_bboxes[ri] if ri < len(data_bboxes) else []
        for ci in range(actual_ncols):
            cell_text = row[ci] if ci < len(row) else ""
            ocr_text = ""
            if ci < len(row_bboxes) and row_bboxes[ci] is not None:
                cx0, cy0, cx1, cy1 = row_bboxes[ci]
                ocr_text = _ocr_region(full_img, cx0, cy0, cx1, cy1, scale)
            final_text = (ocr_text.strip() or cell_text or "").strip()
            cell = t.cell(ri, ci)
            cell.text = final_text
            # 第一列（序号列）居中
            if ci == 0 and final_text.isdigit():
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _calc_col_widths(cell_bboxes: list[list], ncols: int,
                     table_bbox: tuple) -> list:
    """从单元格 bbox 计算列宽（EMU 单位）"""
    # 收集所有列边界
    col_edges = set()
    for row_cells in cell_bboxes:
        for cell in row_cells:
            if isinstance(cell, (list, tuple)) and len(cell) == 4:
                col_edges.add(cell[0])
                col_edges.add(cell[2])
    if len(col_edges) < 2:
        return []

    edges = sorted(col_edges)
    # 取最外层作为起止
    table_x0 = table_bbox[0]
    table_x1 = table_bbox[2]
    table_width_pt = table_x1 - table_x0

    # A4 宽度约 16cm 可用，按比例换算
    doc_width_cm = 16.0
    widths = []
    for i in range(len(edges) - 1):
        w = (edges[i + 1] - edges[i]) / table_width_pt * doc_width_cm
        if w > 0.3:  # 过滤极窄列
            widths.append(Cm(w))
    return widths[:ncols]


# ── 页面处理 ──────────────────────────────────────────────────────────

def _add_page_hybrid(doc: Document, plumber_page, fitz_page, scale: float,
                     full_img: Image.Image) -> None:
    """混合提取：pdfplumber 定位结构 + OCR 提取文字"""
    tables = plumber_page.find_tables()
    table_bboxes = [(t.bbox[1], t.bbox[3], t) for t in tables]
    table_bboxes.sort(key=lambda x: x[0])

    cursor_y = 0.0
    for top_y, bot_y, table in table_bboxes:
        # 表格上方的文本区域
        if top_y > cursor_y + 1:
            text = _ocr_region(full_img, 0, cursor_y, plumber_page.width, top_y, scale)
            if text.strip():
                for line in text.split("\n"):
                    doc.add_paragraph(line)

        _build_table(doc, table, full_img, scale)
        cursor_y = bot_y

    # 页面底部剩余文本
    if cursor_y < plumber_page.height - 1:
        text = _ocr_region(full_img, 0, cursor_y, plumber_page.width, plumber_page.height, scale)
        if text.strip():
            for line in text.split("\n"):
                doc.add_paragraph(line)


def convert_hybrid(pdf_path: Path, docx_path: Path, pages: str | None = None) -> None:
    doc = Document()
    dpi = 300
    scale = dpi / 72

    with fitz.open(str(pdf_path)) as fitz_pdf, pdfplumber.open(str(pdf_path)) as pl_pdf:
        page_indices = _parse_pages(pages) if pages else range(len(fitz_pdf))
        for i, idx in enumerate(page_indices):
            if idx >= len(fitz_pdf):
                continue
            print(f"  处理第 {idx + 1} 页...")
            fitz_page = fitz_pdf[idx]
            plumber_page = pl_pdf.pages[idx]
            full_img = _render_page(fitz_page, dpi=dpi)
            _add_page_hybrid(doc, plumber_page, fitz_page, scale, full_img)
            if i < len(page_indices) - 1:
                doc.add_page_break()
    doc.save(str(docx_path))


# ── 纯文本模式 ────────────────────────────────────────────────────────

def convert_via_text(pdf_path: Path, docx_path: Path, pages: str | None = None) -> None:
    doc = Document()
    with pdfplumber.open(str(pdf_path)) as pdf:
        page_indices = _parse_pages(pages) if pages else range(len(pdf))
        for i, idx in enumerate(page_indices):
            if idx >= len(pdf.pages):
                continue
            _add_page_text(doc, pdf.pages[idx])
            if i < len(page_indices) - 1:
                doc.add_page_break()
    doc.save(str(docx_path))


def _add_page_text(doc: Document, page) -> None:
    tables = page.find_tables()
    table_top_y = [t.bbox[1] for t in tables]
    table_bot_y = [t.bbox[3] for t in tables]

    cursor_y = 0.0
    for i, table in enumerate(tables):
        top = table_top_y[i]
        if top > cursor_y + 1:
            text = page.within_bbox((0, cursor_y, page.width, top)).extract_text()
            if text and (s := text.strip()):
                for line in s.split("\n"):
                    doc.add_paragraph(line)

        rows = table.extract()
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    t.cell(ri, ci).text = (cell or "").strip()

        cursor_y = table_bot_y[i]

    if cursor_y < page.height - 1:
        text = page.within_bbox((0, cursor_y, page.width, page.height)).extract_text()
        if text and (s := text.strip()):
            for line in s.split("\n"):
                doc.add_paragraph(line)


# ── 入口 ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="PDF 转 DOCX 工具")
    parser.add_argument("input", help="PDF 文件路径或目录")
    parser.add_argument("-o", "--output", help="输出 DOCX 路径（仅单文件时有效）")
    parser.add_argument("--pages", help="指定页码范围，如 1-3,5", default=None)
    parser.add_argument("--text", action="store_true", help="纯文本模式（pdfplumber，速度快）")
    args = parser.parse_args()

    source = Path(args.input)
    if not source.exists():
        print(f"错误: {source} 不存在", file=sys.stderr)
        sys.exit(1)

    convert = convert_via_text if args.text else convert_hybrid

    if source.is_file():
        output = Path(args.output) if args.output else source.with_suffix(".docx")
        print(f"转换: {source} -> {output}")
        convert(source, output, args.pages)
        print("完成")
    elif source.is_dir():
        pdfs = sorted(source.glob("*.pdf"))
        if not pdfs:
            print(f"目录 {source} 下没有 PDF 文件", file=sys.stderr)
            sys.exit(1)
        for pdf in pdfs:
            docx = pdf.with_suffix(".docx")
            print(f"转换: {pdf.name} -> {docx.name}")
            convert(pdf, docx, args.pages)
        print(f"批量完成，共 {len(pdfs)} 个文件")
    else:
        print(f"错误: {source} 不是文件或目录", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
