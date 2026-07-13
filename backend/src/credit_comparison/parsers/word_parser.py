from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from src.credit_comparison.core.dto import CompanyProfitLossRecord, FinancialRecord
from src.credit_comparison.core.regex_utils import (
    classify_main_sentence_format,
    detect_calc_scope_hint,
    extract_code_and_name,
    extract_company_detail_section,
    extract_company_detail_items,
    extract_amount_scale,
    extract_direction_amount_unit,
    extract_paraindex,
    extract_quoted_text,
    extract_sheet,
)
PURE_INDEX_LINE_PATTERN = re.compile(r"^\s*[（(]\s*\d+\s*[）)]\s*$")
NUMBER_PLACEHOLDER_PATTERN = re.compile(r"%([1-9])")
logger = logging.getLogger(__name__)


def read_docx_paragraphs(docx_path: str) -> list[str]:
    """读取 docx 段落。"""

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 依赖，无法解析 docx") from exc

    document = Document(docx_path)
    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def format_number_token(value: int, num_fmt: str) -> str:
    """按 Word 编号格式渲染单个序号值。"""

    if num_fmt == "decimal":
        return str(value)
    if num_fmt == "upperLetter":
        return chr(ord("A") + value - 1)
    if num_fmt == "lowerLetter":
        return chr(ord("a") + value - 1)
    if num_fmt in {"upperRoman", "lowerRoman"}:
        roman_pairs = [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]
        rest = value
        tokens: list[str] = []
        for arabic, roman in roman_pairs:
            while rest >= arabic:
                tokens.append(roman)
                rest -= arabic
        result = "".join(tokens)
        return result if num_fmt == "upperRoman" else result.lower()
    return str(value)


def build_numbering_metadata(document: Any) -> dict[str, dict[int, dict[str, Any]]]:
    """从 docx 的 numbering.xml 中提取编号定义。"""

    try:
        from docx.oxml.ns import qn
    except ImportError:
        return {}

    numbering_part = getattr(document.part, "numbering_part", None)
    if numbering_part is None:
        return {}

    numbering_root = numbering_part.element
    num_to_abstract: dict[str, str] = {}
    for num_element in numbering_root.findall(qn("w:num")):
        num_id = str(num_element.get(qn("w:numId")) or "")
        abstract_element = num_element.find(qn("w:abstractNumId"))
        abstract_id = str(abstract_element.get(qn("w:val")) or "") if abstract_element is not None else ""
        if num_id and abstract_id:
            num_to_abstract[num_id] = abstract_id

    abstract_levels: dict[str, dict[int, dict[str, Any]]] = {}
    for abstract_element in numbering_root.findall(qn("w:abstractNum")):
        abstract_id = str(abstract_element.get(qn("w:abstractNumId")) or "")
        if not abstract_id:
            continue
        levels: dict[int, dict[str, Any]] = {}
        for level_element in abstract_element.findall(qn("w:lvl")):
            ilvl = int(level_element.get(qn("w:ilvl")) or 0)
            start_element = level_element.find(qn("w:start"))
            lvl_text_element = level_element.find(qn("w:lvlText"))
            num_fmt_element = level_element.find(qn("w:numFmt"))
            levels[ilvl] = {
                "start": int(start_element.get(qn("w:val")) or 1) if start_element is not None else 1,
                "lvl_text": lvl_text_element.get(qn("w:val")) if lvl_text_element is not None else "（%1）",
                "num_fmt": num_fmt_element.get(qn("w:val")) if num_fmt_element is not None else "decimal",
            }
        abstract_levels[abstract_id] = levels

    metadata: dict[str, dict[int, dict[str, Any]]] = {}
    for num_id, abstract_id in num_to_abstract.items():
        metadata[num_id] = abstract_levels.get(abstract_id, {})
    return metadata


def render_numbering_text(
    lvl_text: str,
    counters: dict[int, int],
    level_definitions: dict[int, dict[str, Any]],
) -> str:
    """按 Word 多级编号模板生成序号文本。"""

    def replace_placeholder(match: re.Match[str]) -> str:
        level_index = int(match.group(1)) - 1
        counter_value = counters.get(level_index)
        if counter_value is None:
            return ""
        number_format = str(level_definitions.get(level_index, {}).get("num_fmt") or "decimal")
        return format_number_token(counter_value, number_format)

    return NUMBER_PLACEHOLDER_PATTERN.sub(replace_placeholder, lvl_text)


def extract_paragraph_numbering_text(
    paragraph: Any,
    numbering_metadata: dict[str, dict[int, dict[str, Any]]],
    numbering_state: dict[str, dict[int, int]],
) -> str:
    """从段落编号属性中恢复序号文本。"""

    try:
        from docx.oxml.ns import qn
    except ImportError:
        return ""

    paragraph_properties = paragraph._p.find(qn("w:pPr"))
    if paragraph_properties is None:
        return ""
    numbering_properties = paragraph_properties.find(qn("w:numPr"))
    if numbering_properties is None:
        return ""

    num_id_element = numbering_properties.find(qn("w:numId"))
    ilvl_element = numbering_properties.find(qn("w:ilvl"))
    num_id = str(num_id_element.get(qn("w:val")) or "") if num_id_element is not None else ""
    if not num_id:
        return ""
    ilvl = int(ilvl_element.get(qn("w:val")) or 0) if ilvl_element is not None else 0

    level_definitions = numbering_metadata.get(num_id, {})
    level_definition = level_definitions.get(ilvl, {})
    start_value = int(level_definition.get("start") or 1)
    current_counters = numbering_state.setdefault(num_id, {})
    current_value = current_counters.get(ilvl, start_value - 1) + 1
    current_counters[ilvl] = current_value
    for level_index in list(current_counters):
        if level_index > ilvl:
            del current_counters[level_index]

    lvl_text = str(level_definition.get("lvl_text") or "（%1）")
    return render_numbering_text(lvl_text, current_counters, level_definitions)


def extract_paragraph_line_texts(paragraph: Any) -> list[str]:
    """提取单个 Word 段落中的逻辑行文本，显式识别手动换行。

    Word 段落内若用 Shift+Enter（`<w:br>`/`<w:cr>`）插入了软换行，
    会把多条指标挤在一个段落里。本函数把它们拆成多个逻辑行，
    供后续逐行解析，避免漏掉指标。
    """

    try:
        from docx.oxml.ns import qn
    except ImportError:
        text = paragraph.text.strip()
        return [text] if text else []

    lines: list[list[str]] = [[]]
    for node in paragraph._p.iter():
        if node.tag in {qn("w:br"), qn("w:cr")}:
            lines.append([])
            continue
        if node.tag == qn("w:t"):
            lines[-1].append(node.text or "")
            continue
        if node.tag == qn("w:tab"):
            lines[-1].append("\t")

    return [text for text in ("".join(parts).strip() for parts in lines) if text]


def read_docx_paragraph_entries(docx_path: str) -> list[dict[str, Any]]:
    """读取 docx 段落，并补回可识别的 Word 编号文本。

    单个段落若含手动换行（`<w:br>`/`<w:cr>`），会产生多条 entry，
    每条带 source_ref（单行 = "{paragraph_index}"，多行 = "{paragraph_index}.{line_index}"），
    编号文本只补到首行。
    """

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 依赖，无法解析 docx") from exc

    document = Document(docx_path)
    numbering_metadata = build_numbering_metadata(document)
    numbering_state: dict[str, dict[int, int]] = {}
    entries: list[dict[str, Any]] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs, 1):
        raw_line_texts = extract_paragraph_line_texts(paragraph)
        numbering_text = extract_paragraph_numbering_text(paragraph, numbering_metadata, numbering_state).strip()
        resolved_line_texts = list(raw_line_texts)
        if numbering_text:
            if resolved_line_texts:
                first_line = resolved_line_texts[0]
                resolved_line_texts[0] = (
                    first_line if first_line.startswith(numbering_text) else f"{numbering_text}{first_line}"
                )
            else:
                resolved_line_texts = [numbering_text]
        if not raw_line_texts and not resolved_line_texts:
            continue

        multi_line = len(resolved_line_texts) > 1
        for line_index, resolved_text in enumerate(resolved_line_texts, start=1):
            raw_text = raw_line_texts[line_index - 1] if line_index - 1 < len(raw_line_texts) else ""
            if not raw_text and not resolved_text:
                continue
            source_ref = f"{paragraph_index}.{line_index}" if multi_line else str(paragraph_index)
            entries.append(
                {
                    "raw_text": raw_text,
                    "resolved_text": resolved_text.strip(),
                    "paragraph_index": paragraph_index,
                    "source_ref": source_ref,
                }
            )
    return entries


def build_candidate_paragraphs(paragraph_entries: list[dict[str, Any]]) -> list[dict[str, str]]:
    """将原始段落预处理为候选块。

    规则如下：
    - 纯序号行 + 下一行指标正文时，合并为一个候选块
    - 已经完整的指标段保持不动
    - 普通说明行、口径约束行保持原样
    """

    candidates: list[dict[str, str]] = []
    index = 0
    while index < len(paragraph_entries):
        current_entry = paragraph_entries[index]
        current = str(current_entry.get("resolved_text") or "").strip()
        current_ref = str(current_entry.get("source_ref") or current_entry.get("paragraph_index") or "")
        if PURE_INDEX_LINE_PATTERN.fullmatch(current) and index + 1 < len(paragraph_entries):
            next_entry = paragraph_entries[index + 1]
            next_paragraph = str(next_entry.get("resolved_text") or "").strip()
            merged = f"{current}{next_paragraph}"
            direction, amount, amount_unit = extract_direction_amount_unit(merged)
            if extract_quoted_text(merged) and direction is not None and amount is not None and amount_unit:
                next_ref = str(next_entry.get("source_ref") or next_entry.get("paragraph_index") or "")
                candidates.append({"text": merged, "source_ref": f"{current_ref}-{next_ref}"})
                index += 2
                continue
        candidates.append({"text": current, "source_ref": current_ref})
        index += 1
    return candidates


def split_by_sheet(paragraphs: list[dict[str, str]]) -> list[dict]:
    """按表单切分段落块。"""

    blocks: list[dict] = []
    current_sheet = ""
    current_paragraphs: list[dict[str, str]] = []

    for paragraph in paragraphs:
        text = str(paragraph.get("text") or "")
        sheet = extract_sheet(text)
        if sheet:
            if current_sheet:
                blocks.append(
                    {
                        "sheet": current_sheet,
                        "paragraphs": current_paragraphs,
                    }
                )
            current_sheet = sheet
            current_paragraphs = []
            continue
        if current_sheet:
            current_paragraphs.append(paragraph)

    if current_sheet:
        blocks.append(
            {
                "sheet": current_sheet,
                "paragraphs": current_paragraphs,
            }
        )
    return blocks


def parse_financial_paragraph(
    paragraph: str,
    file_name: str,
    title: str,
    sheet: str,
    batch_id: str,
    calc_scope_hint: str,
    source_ref: str,
) -> FinancialRecord | None:
    """解析一条指标主记录。"""

    paraindex = extract_paraindex(paragraph)
    quoted_text = extract_quoted_text(paragraph)
    if not quoted_text:
        return None

    code, name = extract_code_and_name(quoted_text)
    direction, amount, amount_unit = extract_direction_amount_unit(paragraph)
    amount_scale = extract_amount_scale(paragraph)
    if not code or not name:
        return None

    record = FinancialRecord(
        title=title,
        sheet=sheet,
        code=code,
        name=name,
        direction=int(direction) if direction is not None else 0,
        amount=amount,
        amount_unit=str(amount_unit or ""),
        amount_scale=int(amount_scale) if amount_scale is not None else 1,
        calc_scope_hint=calc_scope_hint,
        paraindex=paraindex,
        source_ref=source_ref,
        context=paragraph,
        file_name=file_name,
        batch_id=batch_id,
    )
    return record


def parse_company_details(paragraph: str, file_name: str, sheet: str, code: str, batch_id: str) -> list[CompanyProfitLossRecord]:
    """解析一条段落中的企业明细。

    注意：
    - 企业明细只从“主要是/主要为/主要由/主要原因是”之后的文本中提取
    - 不从整段全文直接提取，避免把“本期增加/减少”主句误识别成企业记录
    """

    records: list[CompanyProfitLossRecord] = []
    detail_section = extract_company_detail_section(paragraph)
    if not detail_section:
        return records

    for item in extract_company_detail_items(detail_section):
        record = CompanyProfitLossRecord(
            company=str(item.get("company") or ""),
            direction=int(item.get("direction") or 0),
            profit_loss=item.get("amount"),
            profit_loss_unit=str(item.get("unit") or ""),
            sheet=sheet,
            code=code,
            file_name=file_name,
            batch_id=batch_id,
        )
        setattr(record, "_format_tag", str(item.get("format_tag") or ""))
        setattr(record, "_punctuation_token", str(item.get("punctuation_token") or ""))
        records.append(record)
    return records


def parse_docx_file(docx_path: str, file_name: str, batch_id: str) -> list[dict[str, Any]]:
    """解析 docx 文件并返回主记录及其企业明细。

    这里的策略是“单文件内存聚合”：
    - 单个 Word 文件中的所有 financial_table 记录先放到内存列表
    - 每条主记录携带其解析出的企业明细列表
    - 当前函数只负责解析和聚合，不直接写数据库
    - 数据库写入由上层 service 在单文件解析完成后统一批量执行
    """

    paragraph_entries = read_docx_paragraph_entries(docx_path)
    if not paragraph_entries:
        return []

    raw_paragraphs = [entry["raw_text"] for entry in paragraph_entries if entry.get("raw_text")]
    # for idx, paragraph in enumerate(raw_paragraphs, 1):
    #     logger.info("python-docx段落[%s]: %s", idx, paragraph)

    resolved_paragraphs = [entry["resolved_text"] for entry in paragraph_entries if entry.get("resolved_text")]
    # for idx, paragraph in enumerate(resolved_paragraphs, 1):
    #     logger.info("编号补回段落[%s]: %s", idx, paragraph)

    candidate_paragraphs = build_candidate_paragraphs(paragraph_entries)
    for idx, paragraph in enumerate(candidate_paragraphs, 1):
        logger.info("候选块[%s]: ref=%s, text=%s", idx, paragraph.get("source_ref"), paragraph.get("text"))

    title = str(candidate_paragraphs[0].get("text") or "")
    # 单文件内存聚合，等整个文件解析完成后再由 service 批量写库。
    parsed_records: list[dict[str, Any]] = []

    for block in split_by_sheet(candidate_paragraphs[1:]):
        sheet = block["sheet"]
        current_calc_scope_hint = ""
        for paragraph in block["paragraphs"]:
            paragraph_text = str(paragraph.get("text") or "")
            paragraph_ref = str(paragraph.get("source_ref") or "")
            if extract_paraindex(paragraph_text) is None and extract_quoted_text(paragraph_text) is None:
                scope_hint = detect_calc_scope_hint(paragraph_text)
                if scope_hint:
                    current_calc_scope_hint = scope_hint
                continue
            financial_record = parse_financial_paragraph(
                paragraph_text,
                file_name,
                title,
                sheet,
                batch_id,
                current_calc_scope_hint,
                paragraph_ref,
            )
            if financial_record is None:
                continue
            detail_records = parse_company_details(paragraph_text, file_name, sheet, financial_record.code, batch_id)
            parsed_records.append(
                {
                    "financial": financial_record.to_db_dict(),
                    "company_records": [
                        {
                            **detail.to_db_dict(),
                            "_format_tag": str(getattr(detail, "_format_tag", "") or ""),
                            "_punctuation_token": str(getattr(detail, "_punctuation_token", "") or ""),
                        }
                        for detail in detail_records
                    ],
                }
            )
    return parsed_records


def build_converted_docx_path(converted_dir: str, source_file_path: str) -> str:
    """根据源文件路径生成默认转换后路径。"""

    source_path = Path(source_file_path)
    return str(Path(converted_dir) / f"{source_path.stem}.docx")
