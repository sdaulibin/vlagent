"""
文档比对核心服务

流程：文本提取 → 页级对齐 → 逐页 diff → 结果入库
"""
import json
import time
import logging
from difflib import SequenceMatcher

import pdfplumber
from diff_match_patch import diff_match_patch
from docx import Document
from sqlmodel.ext.asyncio.session import AsyncSession

from src.documents.models import DocumentCompareTask, DocumentPageDiff

logger = logging.getLogger(__name__)

dmp = diff_match_patch()


def extract_pages_from_pdf(file_path: str) -> list[str]:
    """用 pdfplumber 逐页提取原生 PDF 文本"""
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text)
    return pages


def extract_pages_from_docx(file_path: str) -> list[str]:
    """用 python-docx 提取文本，整体视为单页"""
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return ["\n\n".join(paragraphs)]


def extract_pages(file_path: str) -> list[str]:
    """根据文件类型分发提取"""
    lower = file_path.lower()
    if lower.endswith(".pdf"):
        return extract_pages_from_pdf(file_path)
    elif lower.endswith((".docx", ".doc")):
        return extract_pages_from_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {file_path}")


def _page_signature(text: str, max_chars: int = 200) -> str:
    """取页面前 N 个非空白字符作为对齐摘要"""
    stripped = text.replace("\n", " ").replace(" ", "")
    return stripped[:max_chars]


def align_pages(pages_a: list[str], pages_b: list[str]) -> list[tuple]:
    """
    页级对齐：返回 [(page_a_index, page_b_index, diff_type), ...]
    page_a_index / page_b_index 为 0-based，None 表示整页增删
    """
    if not pages_a and not pages_b:
        return []

    # 单页 DOCX：直接对齐不拆分
    if len(pages_a) == 1 and len(pages_b) == 1:
        diff_type = "equal" if pages_a[0] == pages_b[0] else "modified"
        return [(0, 0, diff_type)]

    sigs_a = [_page_signature(p) for p in pages_a]
    sigs_b = [_page_signature(p) for p in pages_b]

    matcher = SequenceMatcher(None, sigs_a, sigs_b)
    aligned = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for i, j in zip(range(i1, i2), range(j1, j2)):
                diff_type = "equal" if pages_a[i] == pages_b[j] else "modified"
                aligned.append((i, j, diff_type))
        elif tag == "replace":
            # 按位置对齐多余的页面
            pairs = min(i2 - i1, j2 - j1)
            for k in range(pairs):
                aligned.append((i1 + k, j1 + k, "modified"))
            # A 多余的页面 = 删除
            for k in range(pairs, i2 - i1):
                aligned.append((i1 + k, None, "deleted"))
            # B 多余的页面 = 新增
            for k in range(pairs, j2 - j1):
                aligned.append((None, j1 + k, "added"))
        elif tag == "delete":
            for i in range(i1, i2):
                aligned.append((i, None, "deleted"))
        elif tag == "insert":
            for j in range(j1, j2):
                aligned.append((None, j, "added"))

    return aligned


def compute_text_diff(text_a: str, text_b: str) -> list[tuple]:
    """
    用 diff-match-patch 计算字符级 diff
    返回 [(op, text), ...]，op: -1=删除, 0=相等, 1=新增
    """
    diffs = dmp.diff_main(text_a, text_b)
    dmp.diff_cleanupSemantic(diffs)
    return [(op, text) for op, text in diffs]


async def process_document_comparison(db: AsyncSession, task: DocumentCompareTask):
    """异步后台任务主函数：提取 → 对齐 → diff → 写库"""
    start_time = time.time()

    try:
        task.status = "processing"
        db.add(task)
        await db.commit()

        # 1. 提取文本
        print(f"[DEBUG] 开始提取文本: {task.file_a_path} vs {task.file_b_path}")
        pages_a = extract_pages(task.file_a_path)
        pages_b = extract_pages(task.file_b_path)

        task.file_a_page_count = len(pages_a)
        task.file_b_page_count = len(pages_b)

        print(f"[DEBUG] 提取完成: A={len(pages_a)}页, B={len(pages_b)}页")
        for i, t in enumerate(pages_a[:5]):
            print(f"[DEBUG]   A page {i+1}: {len(t)} chars, preview: {repr(t[:100])}")
        for i, t in enumerate(pages_b[:5]):
            print(f"[DEBUG]   B page {i+1}: {len(t)} chars, preview: {repr(t[:100])}")

        # 2. 页级对齐
        aligned = align_pages(pages_a, pages_b)

        # 3. 逐页计算 diff 并写入
        for page_a_idx, page_b_idx, diff_type in aligned:
            text_a = pages_a[page_a_idx] if page_a_idx is not None else None
            text_b = pages_b[page_b_idx] if page_b_idx is not None else None

            diff_ops = None
            if diff_type == "modified" and text_a is not None and text_b is not None:
                ops = compute_text_diff(text_a, text_b)
                diff_ops = json.dumps(ops, ensure_ascii=False)

            page_diff = DocumentPageDiff(
                task_id=task.id,
                page_a=(page_a_idx + 1) if page_a_idx is not None else None,
                page_b=(page_b_idx + 1) if page_b_idx is not None else None,
                diff_type=diff_type,
                text_a=text_a,
                text_b=text_b,
                diff_ops_json=diff_ops,
            )
            db.add(page_diff)

        task.comparison_duration = round(time.time() - start_time, 2)
        task.status = "done"
        db.add(task)
        await db.commit()

        logger.info(
            "文档比对完成 task=%d, 页数=%d/%d, 对齐=%d, 耗时=%.2fs",
            task.id, len(pages_a), len(pages_b), len(aligned),
            task.comparison_duration,
        )

    except Exception as e:
        logger.exception("文档比对失败 task=%d", task.id)
        task.status = "failed"
        task.error_msg = str(e)
        task.comparison_duration = round(time.time() - start_time, 2)
        db.add(task)
        await db.commit()
